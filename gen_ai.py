import pathlib
import textwrap
import time
import google.generativeai as genai
from docx import Document
import PyPDF2
from PyPDF2 import PdfReader
from io import BytesIO
import streamlit as st
import time

# Access your secret API key
ak = st.secrets["genai"]["api_key"]

genai.configure(api_key=ak)


for m in genai.list_models():
  if 'generateContent' in m.supported_generation_methods:
    print("models:/n",m.name)


def get_chat_response(chat, prompt):
    # Send the prompt to the chat and get the response
    responses = chat.send_message(prompt, stream=False)

    # Process the responses and concatenate the text
    text_response = []
    for response in responses:
        text_response.append(response.text)

    return "".join(text_response)



def get_contents(chat, raw_texts, model):
    sections = {}
    n=1
    for papers in raw_texts:
        prompt = f'''ok do the literature review for this one, here is the raw_text of the paper:
        {papers}'''
        # Get the rewritten text from the chat
        ind = 'paper' + str(n)
        rewritten_content = get_chat_response(chat, prompt)
        sections[ind] = rewritten_content
        print(model.count_tokens(chat.history))
        print(f"done content for:{ind}")
        print(rewritten_content[:50])
        n = n+1
    return sections

def get_contents_white_paper(chat,section,model):
    sections = {}
    n=1
    
    for heading in section:
        prompt = f'''when giving response  just give the text, The output instructions are dont include any headers like
        'Introduction' or 'Abstract', just give me text. Now give me content for section {heading}  '''
        # Get the rewritten text from the chat
        if n%4==0:
          time.sleep(60)
        rewritten_content = get_chat_response(chat, prompt)
        sections[heading] = rewritten_content
        print(model.count_tokens(chat.history))
        print(f"done content for:{heading}")
        st.write(f"done content for:{heading}")
        print(rewritten_content[:50])
        n = n+1
    msg = "if you feel like any other major headings and content are missing for our white paper summary, generate them."
    rewritten_content = get_chat_response(chat, msg)
    sections["extra"] = rewritten_content

    return sections

def create_new_document(sections, output_path):
        doc = Document()
        for heading, content in sections.items():
            doc.add_heading(heading, level=1)
            doc.add_paragraph(content)
        doc.save(output_path)
        return doc

def send_to_gemini_api(d, chat):
    items = list(d.items())
    print(items)
    length = len(items)
    if length < 4:
        parts = 2
    elif length < 8:
        parts = 4
    elif length < 15:
        parts = 7
    chunk_size = length // parts
    print(chunk_size)
    for i in range(parts):
        start_index = i * chunk_size
        end_index = start_index + chunk_size
        print(start_index, end_index)
        # Create the message for the current chunk
        msg = "\n".join([f"{key}:{value}" for key, value in items[start_index:end_index]])
        print("msg:",msg[:10])
        # Simulate sending the message to the Gemini API
        r = chat.send_message(msg)  # Replace this with your actual API call
        print('res:',r.text[:10])
        print(f"Sent chunk {i+1} to Gemini API")
        print('-' * 20)  # Separator for visual clarity


def init_model_paper_wise(topic, sections, raw_texts):
    print(topic[:3])
    print(sections[:5])
    print(raw_texts[0][:10])
    model=genai.GenerativeModel(
        model_name="gemini-1.5-flash",
        system_instruction=f'''You are a professional researcher who reads and writes research papers and litureature review of research papers.
        I will provide you with several papers on  {topic}\n
        Please conduct a literature review and summarize the key findings, including executive summaries,
        main benefits for employees and organizations, potential drawbacks, implementation strategies, and overall conclusions.'''
        )

    chat = model.start_chat(history=[])

    print(chat)

    msg =f'''Please conduct a literature review on the provided papers focusing on {topic}\n
    For each paper, analyze the paper content and please provide a separate section with the following structure:\n
    {sections}\n 
    Please maintain a unified output format throughout the review and clearly reference specific findings to their corresponding papers.'''
    print(msg)
    r = chat.send_message(msg)
    print(r.text)
    x = get_contents(chat, raw_texts, model)

    output_path = 'paper_wise_summary.docx'
    doc= create_new_document(x, output_path)
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)  # Rewind the buffer to the beginning
    return doc_stream.getvalue()  # Return binary data

def init_model_white_paper(topic, sections, pdf_texts):
    model = genai.GenerativeModel(
    model_name="gemini-1.5-flash",
    system_instruction=f'''
    Imagine you are a reader of research articles. I will give you a document that contains about 10 papers on {topic}
    I want you to analyze every section of that document. Then, I will ask you to generate a white paper summary of those papers. 
    The summary should include the following sections:{sections}
    Generate those sections using the content from the document. Note that the document is created by copying and pasting from different papers, 
    so there might be some numbers, irrelevant symbols, or grammatical errors. Ignore them and focus on generating a coherent and comprehensive summary.
    '''
    )
    chat = model.start_chat(history=[])

    print(chat)

    msg = "when giving response  just give the text, The output instructions are dont include any headers like 'Introduction' or 'Abstract', just give me text. I will give this instruction with evry prompt to remind you "
    r = chat.send_message(msg)
    print(msg, r.text)
    msg = '''Ok, first i will feed you the content of the document by each paper. it is so long so i will feed the content in multiple
    prompts just store the content and analyze it, give me white paper summary sections only when i ask you to ok?'''
    r = chat.send_message(msg)
    print(msg, r.text)
    send_to_gemini_api(pdf_texts, chat)
    msg = '''Ok, I think i feeded you the entrire content, did you get it? just give me confirmation. if you have info,
    Now we will generate my white paper summary, i will specifically ask you to generate each section content like introduction, etc, ok?'''
    r = chat.send_message(msg)
  
    x = get_contents_white_paper(chat,sections,model)

    output_path = 'white_paper_summary.docx'
    
    doc = create_new_document(x, output_path)
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)  # Rewind the buffer to the beginning
    return doc_stream.getvalue()  # Return binary data









if __name__ == "__main__":
   print("gen ai")
